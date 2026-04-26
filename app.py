import time
import re
import io
import streamlit as st
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

# ==========================================
# 1. 유틸리티 함수
# ==========================================

def parse_time_to_seconds(time_str):
    try:
        parts = time_str.split(':')
        return int(parts[0]) * 60 + int(parts[1])
    except: return 999

def set_style(run, font_name="Malgun Gothic", size=13, bold=False, color_rgb=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color_rgb:
        run.font.color.rgb = color_rgb

def clean_script(text):
    text = re.sub(r'#\w+', '', text)
    if "연합뉴스TV 기사문의" in text:
        text = text.split("연합뉴스TV 기사문의")[0]
    return text.strip()

def click_more_button(driver, wait):
    try:
        btn = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "button[class*='button_more']")))
        driver.execute_script("arguments[0].click();", btn)
        time.sleep(0.5)
    except: pass

# ==========================================
# 2. 크롤링 핵심 모듈 (수정됨)
# ==========================================

def get_mbc_anchors_study(driver, doc, target_count=3):
    """ MBC 앵커멘트 수집 (목표: 3개) """
    print("\n--- [1] MBC 앵커멘트 수집 시작 ---")
    driver.get("https://tv.naver.com/imnews?tab=clip")
    time.sleep(3)

    target_links = []
    while len(target_links) < 10:
        items = driver.find_elements(By.CSS_SELECTOR, "a.ClipCardV2_link_thumbnail__NWYf1")
        for item in items:
            if len(target_links) >= 10: break
            try:
                sec = parse_time_to_seconds(item.find_element(By.CSS_SELECTOR, "span.ClipCardV2_playtime__IHYFQ").text)
                link = item.get_attribute("href")
                if 140 <= sec <= 170 and link not in target_links:
                    target_links.append(link)
            except: continue
        if len(target_links) >= 10: break
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)

    wait = WebDriverWait(driver, 5)
    success_count = 0
    for link in target_links:
        if success_count >= target_count: break
        driver.get(link)
        time.sleep(2)
        try:
            click_more_button(driver, wait)
            body = wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "div.ArticleSection_scroll_wrap__ZaUDW"))).text.strip()
            if "◀ 앵커 ▶" not in body: continue

            if "◀ 리포트 ▶" in body: body = body.split("◀ 리포트 ▶")[0]
            if "◀ 기자 ▶" in body: body = body.split("◀ 기자 ▶")[0]

            body = clean_script(re.sub(r'\[.*?\]', '', body).replace("◀ 앵커 ▶", ""))
            success_count += 1
            title_text = f"<앵커멘트{success_count}>"

            p = doc.add_paragraph()
            run_title = p.add_run(title_text + "\n")
            set_style(run_title, size=14, bold=True, color_rgb=RGBColor(0, 112, 192))
            run_body = p.add_run(body)
            set_style(run_body, size=13)
            p.paragraph_format.line_spacing = 1.6
            p.paragraph_format.space_after = Pt(20)
            print(f"✅ {title_text} 완료")
        except: continue

def get_yonhap_shorts_study(driver, doc, target_count=7):
    """ 연합뉴스 단신 수집 (기존 단신+무예독 통합, 목표: 7개) """
    print("\n--- [2] 연합뉴스 단신 수집 시작 ---")
    driver.get("https://tv.naver.com/yonhapnewstv?tab=clip")
    time.sleep(3)

    target_links = []
    while len(target_links) < 15: # 넉넉하게 확보
        items = driver.find_elements(By.CSS_SELECTOR, "a.ClipCardV2_link_thumbnail__NWYf1")
        for item in items:
            if len(target_links) >= 15: break
            try:
                title = item.get_attribute("aria-label")
                sec = parse_time_to_seconds(item.find_element(By.CSS_SELECTOR, "span.ClipCardV2_playtime__IHYFQ").text)
                link = item.get_attribute("href")
                if 40 <= sec <= 55 and "[속보]" not in title and link not in target_links:
                    target_links.append(link)
            except: continue
        if len(target_links) >= 15: break
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)

    wait = WebDriverWait(driver, 5)
    success_count = 0
    for link in target_links:
        if success_count >= target_count: break
        driver.get(link)
        try:
            click_more_button(driver, wait)
            body = clean_script(wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "div.ArticleSection_scroll_wrap__ZaUDW"))).text.strip())
            success_count += 1
            title_text = f"<단신{success_count}>"

            p = doc.add_paragraph()
            run_title = p.add_run(title_text + "\n")
            set_style(run_title, size=14, bold=True, color_rgb=RGBColor(0, 112, 192))
            run_body = p.add_run(body)
            set_style(run_body, size=13)
            p.paragraph_format.line_spacing = 1.6
            p.paragraph_format.space_after = Pt(20)
            print(f"✅ {title_text} 완료")
        except: continue

def get_breaking_news_yonhap(driver, doc, target_count=5):
    """ 실시간 속보 수집 (목표: 5개) """
    print(f"\n--- [3] 실시간 속보 수집 시작 ---")
    driver.get("https://tv.naver.com/yonhapnewstv?tab=clip")
    time.sleep(3)

    breaking_titles = []
    while len(breaking_titles) < target_count:
        items = driver.find_elements(By.CSS_SELECTOR, "a.ClipCardV2_link_thumbnail__NWYf1")
        for item in items:
            if len(breaking_titles) >= target_count: break
            try:
                title = item.get_attribute("aria-label")
                if "[속보]" in title:
                    clean_title = title
                    if "재생시간" in clean_title: clean_title = clean_title.split("재생시간")[0]
                    clean_title = re.sub(r'\s*\d+분\s*\d+초$', '', clean_title).strip()
                    clean_title = re.sub(r'\s*\d+초$', '', clean_title).strip()
                    if clean_title not in breaking_titles:
                        breaking_titles.append(clean_title)
            except: continue
        if len(breaking_titles) >= target_count: break
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)

    p_title = doc.add_paragraph()
    run_title = p_title.add_run("<속보>")
    set_style(run_title, size=14, bold=True, color_rgb=RGBColor(255, 0, 0))
    
    for i, title in enumerate(breaking_titles, 1):
        p_item = doc.add_paragraph()
        run_body = p_item.add_run(f"{i}. {title}")
        set_style(run_body, size=13)
        p_item.paragraph_format.line_spacing = 1.6
        print(f"✅ 속보 {i}번 완료")

# ==========================================
# 3. 메인 컨트롤러
# ==========================================

def main():
    st.set_page_config(page_title="아나운서 원고 자판기", page_icon="🎙️")
    st.title("🎙️ 오늘의 스터디 원고 자판기")
    st.markdown("MBC 앵커멘트(3개), 연합뉴스 단신(7개), 최신 속보(5개)를 한 번에 수집합니다.")

    if st.button("🚀 최신 스터디 원고 생성하기", type="primary"):
        with st.spinner("최신 기사를 수집하고 있습니다. 잠시만 기다려주세요..."):
            driver = None
            try:
                options = Options()
                options.add_argument("--headless=new")
                options.add_argument("--no-sandbox")
                options.add_argument("--disable-dev-shm-usage")
                options.add_argument("--window-size=1920,1080")
                options.add_argument("user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")

                # 로컬 환경 자동 드라이버 설정
                service = Service(ChromeDriverManager().install())
                driver = webdriver.Chrome(service=service, options=options)

                doc = Document()
                for sec in doc.sections:
                    sec.top_margin, sec.bottom_margin = Pt(40), Pt(40)
                    sec.left_margin, sec.right_margin = Pt(50), Pt(50)
                
                # 원고 수집 실행
                get_mbc_anchors_study(driver, doc, target_count=3)
                get_yonhap_shorts_study(driver, doc, target_count=7)
                get_breaking_news_yonhap(driver, doc, target_count=5)
                
                driver.quit()
                
                bio = io.BytesIO()
                doc.save(bio)
                
                st.success("🎉 원고 작성이 완료되었습니다!")
                st.download_button(
                    label="📥 완성된 원고 다운로드 (.docx)",
                    data=bio.getvalue(),
                    file_name=f"Study_Scripts_{time.strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except Exception as e:
                st.error(f"❌ 오류 발생: {e}")
                if driver: driver.quit()

if __name__ == "__main__":
    main()
